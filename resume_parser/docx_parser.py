import io
from docx import Document

def extract_docx_text(file_bytes: bytes) -> str:
    """Extracts text from DOCX bytes using python-docx."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text).strip()
    except Exception as exc:
        raise ValueError(f"Error parsing DOCX: {exc}")
