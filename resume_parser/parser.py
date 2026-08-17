import os
import re
import spacy

from pypdf import PdfReader
from spacy.matcher import PhraseMatcher
from docx import Document
from resume_parser.constants import SKILLS, EDUCATION


# ============================================================
# SPACY SETUP
# ============================================================

nlp = spacy.load("en_core_web_sm")

skill_matcher = PhraseMatcher(
    nlp.vocab,
    attr="LOWER"
)

patterns = [
    nlp.make_doc(skill)
    for skill in SKILLS
]

skill_matcher.add("SKILLS", patterns)


# ============================================================
# PDF TEXT EXTRACTION
# ============================================================

def extract_text_from_pdf(pdf_input):
    """
    Extract text from PDF.

    Supports:
    - file path
    - Streamlit UploadedFile
    """

    reader = PdfReader(pdf_input)

    text = ""

    for page in reader.pages:

        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return text.strip()


# ============================================================
# DOCX TEXT EXTRACTION
# ============================================================

def extract_text_from_docx(docx_input):
    """
    Extract text from DOCX.

    Supports:
    - file path
    - Streamlit UploadedFile
    """

    doc = Document(docx_input)

    text_parts = []

    # Paragraphs
    for paragraph in doc.paragraphs:

        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    # Tables
    for table in doc.tables:

        for row in table.rows:

            row_text = []

            for cell in row.cells:

                if cell.text.strip():
                    row_text.append(cell.text.strip())

            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts).strip()


# ============================================================
# TXT TEXT EXTRACTION
# ============================================================

def extract_text_from_txt(txt_input):
    """
    Extract text from TXT file.

    Supports:
    - Streamlit UploadedFile
    - file path
    """

    if hasattr(txt_input, "getvalue"):

        content = txt_input.getvalue()

        if isinstance(content, bytes):
            return content.decode(
                "utf-8",
                errors="ignore"
            )

        return content

    with open(
        txt_input,
        "r",
        encoding="utf-8",
        errors="ignore"
    ) as file:

        return file.read()


# ============================================================
# GENERIC FILE EXTRACTION
# ============================================================

def extract_text_from_file(file_input):
    """
    Extract text from PDF, DOCX or TXT.
    """

    # Streamlit UploadedFile
    if hasattr(file_input, "name"):

        filename = file_input.name.lower()

    # Normal file path
    else:

        filename = str(file_input).lower()

    if filename.endswith(".pdf"):

        return extract_text_from_pdf(file_input)

    elif filename.endswith(".docx"):

        return extract_text_from_docx(file_input)

    elif filename.endswith(".txt"):

        return extract_text_from_txt(file_input)

    else:

        raise ValueError(
            "Unsupported file format. "
            "Please use PDF, DOCX or TXT."
        )


# ============================================================
# BASIC EMAIL EXTRACTION
# ============================================================

def extract_email(text):

    pattern = (
        r"[a-zA-Z0-9._%+-]+"
        r"@[a-zA-Z0-9.-]+\."
        r"[a-zA-Z]{2,}"
    )

    match = re.search(
        pattern,
        text
    )

    return match.group() if match else None


# ============================================================
# BASIC PHONE EXTRACTION
# ============================================================

def extract_phone(text):

    pattern = r"(\+91[\-\s]?)?[6-9]\d{9}"

    match = re.search(
        pattern,
        text
    )

    return match.group() if match else None


# ============================================================
# BASIC SKILL EXTRACTION
# ============================================================

def extract_skills(text):

    doc = nlp(text)

    matches = skill_matcher(doc)

    skills = set()

    for _, start, end in matches:

        skills.add(
            doc[start:end].text
        )

    return sorted(skills)


# ============================================================
# BASIC EDUCATION EXTRACTION
# ============================================================

def extract_education(text):

    found = []

    text_lower = text.lower()

    for degree in EDUCATION:

        if degree.lower() in text_lower:

            found.append(degree)

    return found


# ============================================================
# BASIC COLLEGE EXTRACTION
# ============================================================

def extract_college(text):

    keywords = [
        "college",
        "university",
        "institute"
    ]

    for line in text.splitlines():

        clean_line = line.strip()

        if not clean_line:
            continue

        for keyword in keywords:

            if keyword in clean_line.lower():

                return clean_line

    return None


# ============================================================
# BASIC EXPERIENCE EXTRACTION
# ============================================================

def extract_experience(text):

    pattern = (
        r"(\d+(?:\.\d+)?)"
        r"\+?\s*"
        r"(year|years|yr|yrs)"
    )

    match = re.search(
        pattern,
        text,
        re.IGNORECASE
    )

    if match:

        return match.group()

    return "Fresher"


# ============================================================
# BASIC PROJECT EXTRACTION
# ============================================================

def extract_projects(text):

    projects = []

    lines = text.splitlines()

    capture = False

    stop_words = [
        "education",
        "experience",
        "skills",
        "certifications",
        "achievement",
        "achievements",
        "languages",
        "interests"
    ]

    for line in lines:

        clean = line.strip()

        if not clean:
            continue

        if clean.lower() in [
            "projects",
            "project"
        ]:

            capture = True
            continue

        if capture:

            if clean.lower() in stop_words:
                break

            projects.append(clean)

    return projects


# ============================================================
# BASIC PARSER
# ============================================================

def parse_resume_basic(file_input):

    text = extract_text_from_file(
        file_input
    )

    if not text.strip():

        raise ValueError(
            "Could not extract text from resume."
        )

    data = {

        "text": text,

        "email": extract_email(text),

        "phone": extract_phone(text),

        "skills": extract_skills(text),

        "education": extract_education(text),

        "college": extract_college(text),

        "experience": extract_experience(text),

        "projects": extract_projects(text)
    }

    return data


# ============================================================
# RESUME PATH
# ============================================================

def get_resume_path(folder="Data"):

    base_dir = os.path.dirname(
        os.path.dirname(
            os.path.abspath(__file__)
        )
    )

    folder_path = os.path.join(
        base_dir,
        folder
    )

    if not os.path.exists(folder_path):

        return None

    for file in os.listdir(folder_path):

        if file.lower().endswith(
            (".pdf", ".docx", ".txt")
        ):

            return os.path.join(
                folder_path,
                file
            )

    return None


# ============================================================
# TEST
# ============================================================

if __name__ == "__main__":

    resume_path = get_resume_path()

    if resume_path is None:

        print("No Resume Found")

    else:

        result = parse_resume_basic(
            resume_path
        )

        print(
            "\n========== BASIC RESUME PARSER ==========\n"
        )

        for key, value in result.items():

            if key != "text":

                print(
                    f"{key.upper()} : {value}"
                )