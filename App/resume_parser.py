import os
import re
import json
import spacy
import fitz
import io
import zipfile
import xml.etree.ElementTree as ET
from pypdf import PdfReader
from docx import Document
from spacy.matcher import PhraseMatcher
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from constants import SKILLS, DEGREE_PATTERNS, SECTION_ALIASES



# def extract_text_from_pdf(pdf_input):
#     pdf_input.seek(0)
    
#     file_bytes = pdf_input.read()
    
#     document = fitz.open(stream=file_bytes, filetype="pdf")
    
#     text = ""
    
#     for page in document:
#         text += page.get_text("text") + "\n"
        
#     document.close()
#     return text
load_dotenv()

llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    temperature=0
)

def extract_text_from_pdf(pdf_input):
    pdf_input.seek(0)

    file_bytes = pdf_input.read()

    document = fitz.open(stream=file_bytes, filetype="pdf")

    text = ""

    for page in document:
        blocks = page.get_text("blocks")  # (x0, y0, x1, y1, text, block_no, block_type)

        # Keep only real text blocks, sort top-to-bottom, left-to-right
        text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]
        text_blocks.sort(key=lambda b: (round(b[1], 1), b[0]))

        for block in text_blocks:
            text += block[4].strip() + "\n"

    document.close()
    return text


def extract_text_from_docx(docx_input):

    # Reset file position
    if hasattr(docx_input, "seek"):
        docx_input.seek(0)

    # Read DOCX into memory
    file_bytes = docx_input.read()

    # -------------------------------
    # 1. Normal paragraphs + tables
    # -------------------------------
    document = Document(io.BytesIO(file_bytes))

    text_parts = []

    # Normal paragraphs
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            text_parts.append(text)

    # Tables
    for table in document.tables:

        for row in table.rows:

            row_text = []

            for cell in row.cells:

                cell_text = cell.text.strip()

                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                text_parts.append(" | ".join(row_text))

    # -------------------------------
    # 2. Extract text from headers,
    #    footers and text boxes
    # -------------------------------

    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    }

    try:

        with zipfile.ZipFile(io.BytesIO(file_bytes)) as docx_zip:

            # XML files that can contain resume text
            xml_files = [
                name for name in docx_zip.namelist()
                if (
                    name == "word/document.xml"
                    or name.startswith("word/header")
                    or name.startswith("word/footer")
                    or name == "word/footnotes.xml"
                    or name == "word/endnotes.xml"
                )
            ]

            for xml_file in xml_files:

                xml_data = docx_zip.read(xml_file)

                root = ET.fromstring(xml_data)

                # Text inside text boxes
                for textbox in root.findall(
                    ".//w:txbxContent",
                    namespaces
                ):

                    textbox_text = []

                    for node in textbox.findall(
                        ".//w:t",
                        namespaces
                    ):

                        if node.text:
                            textbox_text.append(node.text)

                    if textbox_text:

                        text = " ".join(textbox_text).strip()

                        if text:
                            text_parts.append(text)

                # Header/footer text
                if (
                    xml_file.startswith("word/header")
                    or xml_file.startswith("word/footer")
                ):

                    header_footer_text = []

                    for node in root.findall(
                        ".//w:t",
                        namespaces
                    ):

                        if node.text:
                            header_footer_text.append(node.text)

                    if header_footer_text:

                        text = " ".join(header_footer_text).strip()

                        if text:
                            text_parts.append(text)

    except Exception as e:

        print("DOCX XML extraction warning:", e)

    # -------------------------------
    # 3. Remove duplicate lines
    # -------------------------------

    final_text = []

    seen = set()

    for text in text_parts:

        text = re.sub(r"\s+", " ", text).strip()

        if text and text.lower() not in seen:

            final_text.append(text)

            seen.add(text.lower())

    return "\n".join(final_text)

def extract_text_from_txt(txt_input):
    txt_input.seek(0)
    return txt_input.read().decode("utf-8", errors="ignore")


def extract_text_from_file(file_input):
    filename = file_input.name.lower()
    
    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_input)
    
    elif filename.endswith(".docx"):
            return extract_text_from_docx(file_input)
    
    elif filename.endswith(".txt"):
            return extract_text_from_txt(file_input)
        
    else :
        raise ValueError(
            "Unsupported file format. Please upload PDF, DOCX, TXT."
        )
    
    
def clean_text(text):
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()

def parse_resume_with_gemini(text):
    
    prompt = f"""
You are an expert resume parser.

Analyze the resume and extract candidate information.

Return ONLY valid JSON.

Use exactly this structure:

{{
    "name": null,
    "email": null,
    "phone": null,
    "college": null,
    "education": [],
    "experience": null,
    "skills": [],
    "projects": [],
    "hobbies": []
}}

IMPORTANT RULES:

1. Do not invent information.
2. Extract only information actually present in the resume.
3. "education" MUST be a list of strings.
4. "skills" MUST be a list of strings.
5. "projects" MUST be a list of strings.
6. "hobbies" MUST be a list of strings.
7. "experience" MUST be a single string.
8. "college" MUST be a single string.
9. If information is missing, use null, [] or "Fresher".
10. Return valid JSON only.

RESUME:

{text}
"""

    response = llm.invoke(prompt)

    result = response.content.strip()

    # Remove markdown code fences if Gemini returns them
    if result.startswith("```"):
        result = result.replace("```json", "")
        result = result.replace("```", "")
        result = result.strip()

    try:
        return json.loads(result)

    except json.JSONDecodeError:
        return {
            "error": "Gemini returned invalid JSON",
            "raw_response": result
        }

def get_lines(text):
    return [line.strip() for line in text.splitlines() if line.strip()]


def extract_email(text):
    pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"

    match = re.search(pattern, text)

    return match.group(0) if match else None


def extract_phone(text):
    patterns = [
        r"\+91[\s-]?[6-9]\d{9}",
        r"\b[6-9]\d{9}\b",
        r"\+1[\s-]?\d{3}[\s-]?\d{3}[\s-]?\d{4}",
        r"\b\d{7,15}\b"
    ]

    for pattern in patterns:
        match = re.search(pattern, text)

        if match:
            return match.group(0)

    return None


def normalize_heading(line):
    line = line.lower().strip()

    line = re.sub(r"[^a-zA-Z\s]", "", line)

    line = re.sub(r"\s+", " ", line)

    return line.strip()


def looks_like_heading(line):
    words = line.split()

    if not (1 <= len(words) <= 6):
        return False

    if re.search(r"[.,]", line):
        return False

    if line.isupper():
        return True

    if all(word[0].isupper() for word in words if word[0].isalpha()):
        return True

    return False


def detect_sections(lines):
    sections = {}

    current_section = None

    for line in lines:

        normalized = normalize_heading(line)

        detected_section = None

        for section, aliases in SECTION_ALIASES.items():

            if normalized in aliases:
                detected_section = section
                break

        if detected_section:

            current_section = detected_section

            if current_section not in sections:
                sections[current_section] = []

            continue

        if current_section:
            sections[current_section].append(line)

    return sections


def extract_name(text):
    lines = get_lines(text)

    candidates = []

    rejected = [
        "resume",
        "curriculum",
        "profile",
        "summary",
        "objective",
        "skills",
        "education",
        "experience",
        "projects",
        "certification",
        "student",
        "engineer",
        "engineering"
    ]

    # Only inspect the beginning of the resume
    for index, line in enumerate(lines[:15]):

        # Don't consider contact-information lines
        if "@" in line:
            continue

        if re.search(r"\d{5,}", line):
            continue

        words = line.split()

        # A normal name is usually 2-5 words
        if not 2 <= len(words) <= 5:
            continue

        lower_line = line.lower()

        # Reject obvious headings
        if any(word in lower_line for word in rejected):
            continue

        # Name should contain alphabetic words
        if not all(
            re.fullmatch(r"[A-Za-z][A-Za-z'.-]*", word)
            for word in words
        ):
            continue

        score = 0

        # Earlier lines get higher score
        score += max(0, 15 - index)

        # Prefer 2-4 word names
        if 2 <= len(words) <= 4:
            score += 5

        # Prefer capitalized words
        score += sum(
            2 for word in words
            if word[0].isupper()
        )

        # Strong signal if entire line is uppercase
        if line.isupper():
            score += 5

        candidates.append((score, line))

    if candidates:
        candidates.sort(reverse=True)

        return candidates[0][1].title()

    return None

def extract_education(text):
    lines = get_lines(text)

    education = []
    seen = set()

    candidates = list(lines) + [
        f"{lines[i]} {lines[i + 1]}"
        for i in range(len(lines) - 1)
    ]

    for line in candidates:

        if line in seen:
            continue

        lower = line.lower()

        for degree in DEGREE_PATTERNS:

            pattern = r"(?<!\w)" + re.escape(degree.lower()) + r"(?!\w)"

            if re.search(pattern, lower):
                education.append(line)
                seen.add(line)
                break

    return education


def extract_college(text):
    lines = get_lines(text)

    keywords = [
        "college",
        "university",
        "institute",
        "school",
        "iit",
        "nit",
        "bits"
    ]

    for line in lines:

        lower = line.lower()

        if any(keyword in lower for keyword in keywords):

            if "college student" in lower:
                continue

            return line

    return None

def extract_experience(text):

    patterns = [
        r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience|exp)?",
        r"(?:experience|exp)\s*(?:of\s+)?(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)",
        r"(\d+(?:\.\d+)?)\+?\s*(?:months?|mos?)"
    ]

    for pattern in patterns:

        match = re.search(
            pattern,
            text,
            flags=re.IGNORECASE
        )

        if match:
            return match.group(0)

    return "Fresher"


def extract_skills(text):

    text_lower = text.lower()

    found_skills = []

    for skill in SKILLS:

        skill_lower = skill.lower()

        pattern = (
            r"(?<!\w)"
            + re.escape(skill_lower)
            + r"(?!\w)"
        )

        if re.search(pattern, text_lower):

            found_skills.append(skill)

    return sorted(set(found_skills))


def extract_projects(text):

    lines = get_lines(text)

    sections = detect_sections(lines)

    project_lines = sections.get("projects", [])

    if not project_lines:
        return []

    def ends_with_terminal_punctuation(s):
        return s.rstrip().endswith((".", "!", "?"))

    projects = []
    current_project = None

    for line in project_lines:

        is_new_title = False

        if current_project is None:
            is_new_title = True
        else:
            desc = current_project["description"]
            if desc and ends_with_terminal_punctuation(desc) and not ends_with_terminal_punctuation(line):
                is_new_title = True

        if is_new_title:

            if current_project:
                projects.append(current_project)

            current_project = {
                "name": line,
                "description": ""
            }

        else:

            if current_project["description"]:
                current_project["description"] += " "

            current_project["description"] += line

    if current_project:
        projects.append(current_project)

    return projects


def extract_hobbies_interests(text):

    lines = get_lines(text)

    sections = detect_sections(lines)

    hobbies_section = sections.get("hobbies", [])

    hobbies = []

    for line in hobbies_section:

        if line.strip():
            hobbies.append(line.strip())

    return hobbies


def parse_resume(file_input):

    # 1. Extract text from PDF/DOCX/TXT
    text = extract_text_from_file(file_input)

    # 2. Clean extracted text
    text = clean_text(text)

    # 3. Send resume text to Gemini
    data = parse_resume_with_gemini(text)

    return data


def get_resume_path(folder="Data"):
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    folder_path = os.path.join(base_dir, folder)

    if not os.path.exists(folder_path):
        return None

    for file in os.listdir(folder_path):
        if file.lower().endswith(".pdf"):
            return os.path.join(folder_path, file)

    return None


if __name__ == "__main__":

    resume_path = get_resume_path()

    if resume_path is None:
        print("No Resume Found")

    else:
        candidate = parse_resume(resume_path)

        print("\n========== PARSED RESUME ==========\n")

        for key, value in candidate.items():
            print(f"{key.upper()} : {value}")
            

def extract_text_from_txt(txt_input):
    """Extract text from TXT file"""
    try:
        return txt_input.getvalue().decode('utf-8')
    except Exception as e:
        return f"Error extracting TXT: {str(e)}"


def extract_text_from_file(file_input):
    """Extract text from any file format (PDF, DOCX, TXT)."""

    filename = file_input.name.lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_input)

    elif filename.endswith(".docx"):
        return extract_text_from_docx(file_input)

    elif filename.endswith(".txt"):
        return extract_text_from_txt(file_input)

    else:
        return "Error: Unsupported file format. Please use PDF, DOCX, or TXT."